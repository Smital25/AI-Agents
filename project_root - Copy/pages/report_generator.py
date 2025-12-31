"""
📄 TeamCollab — IEEE Report AI Agent — DeepSeek (Ollama, Auto Mongo Integration)

Features:
- Auto-reads for ACTIVE PROJECT (from st.session_state["active_project"]):
    • Latest literature review (research_reviews)
    • Final objectives (project_objectives)
    • Tasks / modules (tasks)
- Generates full IEEE-style report (single-shot OR multi-chapter)
- Auto-saves report versions into MongoDB (project_reports)
- Grammar polishing (LLM-based)
- References generator (LLM-based)
- Heuristic plagiarism scoring (self-similarity only)
- AI Evaluation & Improvement suggestions
- Export to TXT, MD, DOCX, PDF with IEEE-like styling

Requirements:
- Ollama running locally with: deepseek-r1:32b  (or adjust env IEEE_OLLAMA_MODEL / OLLAMA_MODEL)
- MongoDB access
"""

from __future__ import annotations

import io
import re
import os
import datetime
from typing import Optional, Dict, Any, List

import requests
import streamlit as st
from pymongo import MongoClient

# Optional exports
try:
    from docx import Document  # python-docx
except Exception:
    Document = None

try:
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.enums import TA_JUSTIFY
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
except Exception:
    SimpleDocTemplate = None
    Paragraph = None
    Spacer = None
    getSampleStyleSheet = None
    ParagraphStyle = None
    A4 = None
    TA_JUSTIFY = None
    pdfmetrics = None
    TTFont = None


# ---------------- CONFIG ----------------

# Ollama / DeepSeek
OLLAMA_URL = os.getenv("OLLAMA_URL", "http://127.0.0.1:11434/api/chat")
OLLAMA_MODEL = os.getenv("IEEE_OLLAMA_MODEL", os.getenv("OLLAMA_MODEL", "deepseek-r1:32b"))
OLLAMA_TIMEOUT = int(os.getenv("OLLAMA_TIMEOUT", "300"))

# MongoDB
MONGO_URI = os.getenv("MONGO_URI", "mongodb://localhost:27017/")
client = MongoClient(MONGO_URI)
db = client["team_collab_db"]

research_col = db["research_reviews"]
objectives_col = db["project_objectives"]
tasks_col = db["tasks"]
reports_col = db["project_reports"]

st.set_page_config(
    page_title="IEEE Report Agent (DeepSeek)",
    page_icon="📄",
    layout="centered",
)

# ---------------- SMALL HELPERS ----------------

def safe_get(d: Any, k: str, default: Any = "") -> Any:
    return d.get(k, default) if isinstance(d, dict) else default


# ---------------- OLLAMA CALL ----------------

def call_llm(prompt: str, max_tokens: int = 6000) -> str:
    """
    Auto-select best working ollama DeepSeek/LLM model.
    Falls back smoothly.
    """
    CANDIDATE_MODELS = [
        OLLAMA_MODEL,                    # your env
        "deepseek-r1:32b",
        "deepseek-r1:14b",
        "deepseek-r1:7b",
        "deepseek-r1",
        "llama3.2",
        "mistral",
    ]

    for m in CANDIDATE_MODELS:
        try:
            r = requests.post(
                OLLAMA_URL,
                json={
                    "model": m,
                    "messages": [{"role": "user", "content": prompt}],
                    "stream": False,
                    "temperature": 0.2,
                },
                timeout=OLLAMA_TIMEOUT,
            )
            if r.status_code != 200:
                continue

            j = r.json()
            result = (j.get("message", {}) or {}).get("content", "")
            if result.strip():
                return result
        except Exception:
            continue

    return "Sorry, no local LLM model responded. Please start ollama serve."


# ---------------- PROMPTS ----------------

IEEE_SECTIONS = [
    "Abstract",
    "Introduction",
    "Literature Review",
    "Problem Definition",
    "Objectives",
    "Methodology",
    "Proposed System",
    "System Architecture",
    "Implementation",
    "Results and Discussion",
    "Conclusion",
    "Future Scope",
    "References",
]


def build_ieee_prompt(
    project_title: str,
    domain: str,
    objectives: str,
    literature: str,
    implementation: str,
) -> str:
    return f"""
You are an experienced academic advisor and IEEE journal author.

Write a COMPLETE IEEE-style project report in clear paragraphs,
assuming this is a B.E / B.Tech / M.Tech level academic project.

PROJECT TITLE: {project_title}
DOMAIN: {domain}

PROJECT OBJECTIVES:
{objectives or "No objectives explicitly stored; infer reasonable ones from context."}

LITERATURE REVIEW INPUT (raw notes or summary):
{literature or "No stored literature; use domain knowledge to write a correct academic review."}

IMPLEMENTATION DETAILS (modules / stack / flow):
{implementation or "Implementation details not fully stored; assume a reasonable stack and modules."}

STRUCTURE & ORDER (use these headings EXACTLY in this order):
1. Abstract
2. Introduction
3. Literature Review
4. Problem Definition
5. Objectives
6. Methodology
7. Proposed System
8. System Architecture
9. Implementation
10. Results and Discussion
11. Conclusion
12. Future Scope
13. References

STYLE RULES:
- Use formal academic tone.
- Write long-form paragraphs (no bullet points except if really necessary).
- Where suitable, mention technologies, frameworks, and algorithms.
- In "References", create 6–10 realistic academic-style references (IEEE-like).
- DO NOT mention that you are an AI.
- DO NOT add extra headings apart from those listed.
"""


def build_section_prompt(
    section: str,
    project_title: str,
    domain: str,
    objectives: str,
    literature: str,
    implementation: str,
) -> str:
    """
    Used for multi-chapter generation: one section at a time.
    """
    base_context = f"""
Project Title: {project_title}
Domain: {domain}

Objectives:
{objectives or "Not explicitly stored; derive reasonable objectives."}

Literature Context:
{literature or "Not stored; use plausible academic context for this domain."}

Implementation / Tasks:
{implementation or "Only partial implementation data available; assume a reasonable design."}
"""
    return f"""
You are an IEEE-style academic writer.

Using ONLY the context below, write the **{section}** section
of a full IEEE project report.

CONTEXT:
\"\"\"{base_context[:10000]}\"\"\"


SECTION TO WRITE: {section}

REQUIREMENTS:
- Do NOT write any other sections, only "{section}".
- Use a formal academic tone.
- 2–6 paragraphs (depending on section).
- No bullet lists unless clearly necessary.
- For "References", output 6–10 IEEE-style references [1], [2], etc.
- Do NOT mention that you are an AI.

Now write the {section} section:
"""


def build_grammar_prompt(text: str) -> str:
    return f"""
You are an academic English editor.

Task:
- Correct grammar, spelling, punctuation.
- Improve clarity and flow.
- Keep the meaning and structure of sections.
- Preserve headings exactly (Abstract, Introduction, etc.) if present.
- Do NOT shorten too much. Just refine.

Text:
\"\"\"{text}\"\"\"


Return ONLY the improved text.
"""


def build_references_prompt(text: str, project_title: str, domain: str) -> str:
    return f"""
You are an academic writing assistant.

From the following IEEE-style report text, infer suitable academic references.
You may invent realistic titles, authors, venues, and years, but they must logically match
the topic.

PROJECT TITLE: {project_title}
DOMAIN: {domain}

REPORT TEXT (INCOMPLETE, may not have references):
\"\"\"{text[:8000]}\"\"\"


TASK:
- Generate 6 to 12 IEEE-style references.
- Use numbering: [1] ..., [2] ..., etc.
- Mix journals, conferences, and relevant websites.
- Do NOT explain anything, just list the references.
"""


def build_eval_prompt(text: str, project_title: str, domain: str) -> str:
    return f"""
You are an expert academic evaluator and supervisor.

Evaluate the following IEEE-style report and suggest improvements.

PROJECT TITLE: {project_title}
DOMAIN: {domain}

REPORT:
\"\"\"{text[:20000]}\"\"\"


TASK:
Provide a structured evaluation with the following headings:

1. Overall Score (0–10)
2. Strengths
3. Weaknesses / Gaps
4. IEEE Structural Compliance
5. Technical Depth
6. Suggestions for Improvement (detailed)

Use plain text. Do NOT repeat the whole report. Do NOT output JSON.
"""


# ---------------- HEURISTIC "PLAGIARISM" SCORE ----------------

def heuristic_plagiarism_score(text: str) -> float:
    """
    This is NOT real plagiarism detection.
    It only checks self-similarity & repetition inside the given text.

    High repetition => higher "risk" score.
    Return 0–100.
    """
    cleaned = re.sub(r"\s+", " ", text or "").strip()
    if not cleaned:
        return 0.0

    # Sentence-based uniqueness
    sentences = re.split(r'(?<=[.!?])\s+', cleaned)
    sentences = [s.strip() for s in sentences if s.strip()]
    total_sents = len(sentences)

    if total_sents == 0:
        return 0.0

    unique_sents = len(set(sentences))

    # n-gram repetition
    words = cleaned.lower().split()
    n = 4
    if len(words) < n:
        uniq_ratio = unique_sents / total_sents
        risk = (1.0 - uniq_ratio) * 100.0
        return round(risk, 2)

    ngrams = []
    for i in range(len(words) - n + 1):
        ngrams.append(" ".join(words[i:i + n]))

    total_ngrams = len(ngrams)
    unique_ngrams = len(set(ngrams))

    sent_uniqueness = unique_sents / total_sents
    ngram_uniqueness = unique_ngrams / max(1, total_ngrams)

    # lower uniqueness => higher risk
    avg_uniqueness = (sent_uniqueness + ngram_uniqueness) / 2.0
    risk = (1.0 - avg_uniqueness) * 100.0
    risk = max(0.0, min(100.0, risk))
    return round(risk, 2)


def explain_plagiarism_score(score: float) -> str:
    if score < 20:
        return "Low repetition / high internal originality (heuristic only)."
    if score < 40:
        return "Some repetition but overall reasonably varied."
    if score < 60:
        return "Moderate repetition; consider rephrasing some parts."
    if score < 80:
        return "High repetition; structure and wording may be too similar across sections."
    return "Very high repetition; large parts look internally duplicated. Strong rewrite recommended."


# ---------------- FETCH FROM MONGO ----------------

def fetch_project_literature(pid: str) -> str:
    doc = research_col.find_one({"project_id": pid}, sort=[("created_at", -1)])
    if not doc:
        return ""
    pieces: List[str] = []
    for p in doc.get("results", []) or []:
        title = safe_get(p, "title", "")
        source = safe_get(p, "source", "")
        year = safe_get(p, "year", "")
        summary = safe_get(p, "summary", "")
        abstract = safe_get(p, "abstract", "")
        pieces.append(f"{title} ({source}, {year}) — {summary} {abstract}")
    return "\n".join(pieces)


def fetch_project_objectives(pid: str) -> str:
    doc = objectives_col.find_one({"project_id": pid})
    if not doc:
        return ""
    objs = doc.get("objectives", []) or []
    return "\n".join(f"{i+1}. {o}" for i, o in enumerate(objs) if o)


def fetch_project_implementation(pid: str) -> str:
    items = list(tasks_col.find({"project_id": pid}))
    if not items:
        return ""
    lines: List[str] = []
    for t in items:
        name = safe_get(t, "task_name", safe_get(t, "title", ""))
        module = safe_get(t, "module", "")
        status = safe_get(t, "status", "")
        days = safe_get(t, "est_days", safe_get(t, "estimated_days", ""))
        desc = safe_get(t, "description", "")
        lines.append(
            f"Task: {name}, Module: {module}, Status: {status}, EstDays: {days}, Details: {desc}"
        )
    return "\n".join(lines)


def save_report_version(
    pid: str,
    project_title: str,
    domain: str,
    text: str,
    auto: bool = True,
) -> None:
    if not pid or not text.strip():
        return
    reports_col.insert_one(
        {
            "project_id": pid,
            "project_name": project_title,
            "domain": domain,
            "text": text,
            "auto_saved": auto,
            "created_at": datetime.datetime.utcnow(),
        }
    )


def get_last_report(pid: str) -> Optional[Dict[str, Any]]:
    if not pid:
        return None
    return reports_col.find_one({"project_id": pid}, sort=[("created_at", -1)])


# ---------------- EXPORT HELPERS ----------------

def export_docx(text: str, project_title: str) -> Optional[io.BytesIO]:
    if Document is None:
        return None
    doc = Document()
    for line in (text or "").splitlines():
        if not line.strip():
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def register_pdf_font() -> str:
    # For Unicode; if fail, fallback to Helvetica
    if pdfmetrics is None or TTFont is None:
        return "Helvetica"
    try:
        font_paths = [
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        ]
        for path in font_paths:
            if os.path.exists(path):
                pdfmetrics.registerFont(TTFont("DejaVuSans", path))
                return "DejaVuSans"
    except Exception:
        pass
    return "Helvetica"


PDF_BASE_FONT = register_pdf_font()


def export_pdf(text: str, project_title: str) -> Optional[io.BytesIO]:
    """
    Export PDF using ReportLab. Returns BytesIO or None.
    Safe against missing imports and Pylance optional warnings.
    """
    # ---- Import inside to guarantee non-None for Pylance ----
    try:
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.enums import TA_JUSTIFY
    except Exception:
        return None

    buf = io.BytesIO()

    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=40,
        leftMargin=40,
        topMargin=40,
        bottomMargin=40,
        title=project_title,
    )

    styles = getSampleStyleSheet()
    normal_style = styles.get("Normal")

    # fallback for TA_JUSTIFY
    try:
        justify = TA_JUSTIFY
    except Exception:
        justify = 4  # integer code for justify

    body_style = ParagraphStyle(
        name="Body",
        parent=normal_style,
        fontName=PDF_BASE_FONT,
        fontSize=10,
        alignment=justify,
        leading=14,
    )

    title_style = styles.get("Title")

    elems = []
    elems.append(Paragraph(project_title or "IEEE Report", title_style))
    elems.append(Spacer(1, 12))

    # split by double newline
    blocks = re.split(r"\n\s*\n", (text or "").strip())
    for blk in blocks:
        blk = blk.strip()
        if not blk:
            continue

        # basic section heading guess
        lines = blk.splitlines()
        if len(lines) == 1 and ":" not in lines[0] and len(lines[0].split()) <= 6:
            elems.append(Paragraph(f"<b>{lines[0]}</b>", body_style))
            elems.append(Spacer(1, 8))
        else:
            safe = blk.replace("\n", " ")
            elems.append(Paragraph(safe, body_style))
            elems.append(Spacer(1, 6))

    doc.build(elems)
    buf.seek(0)
    return buf



# ---------------- GENERATION LOGIC ----------------

def generate_ieee_single(
    project_title: str,
    domain: str,
    objectives: str,
    literature: str,
    implementation: str,
) -> str:
    prompt = build_ieee_prompt(project_title, domain, objectives, literature, implementation)
    text = call_llm(prompt)
    return text.strip() if text else ""


def generate_ieee_multi(
    project_title: str,
    domain: str,
    objectives: str,
    literature: str,
    implementation: str,
) -> str:
    """
    Multi-chapter generation: calls DeepSeek once per IEEE section
    and concatenates the sections.
    """
    sections_out: List[str] = []
    for sec in IEEE_SECTIONS:
        sec_prompt = build_section_prompt(
            sec,
            project_title,
            domain,
            objectives,
            literature,
            implementation,
        )
        sec_text = call_llm(sec_prompt, max_tokens=1800) or ""
        sec_text = sec_text.strip()
        if not sec_text:
            continue
        sections_out.append(f"{sec}\n\n{sec_text}")
    return "\n\n".join(sections_out).strip()


def polish_text(text: str) -> str:
    if not text.strip():
        return ""
    out = call_llm(build_grammar_prompt(text))
    return out.strip() if out else ""


def generate_references(text: str, project_title: str, domain: str) -> str:
    if not text.strip():
        return ""
    out = call_llm(build_references_prompt(text, project_title, domain))
    return out.strip() if out else ""


def evaluate_report(text: str, project_title: str, domain: str) -> str:
    if not text.strip():
        return ""
    out = call_llm(build_eval_prompt(text, project_title, domain), max_tokens=1500)
    return out.strip() if out else ""


# ---------------- MAIN PAGE ----------------

def ieee_report_agent_page():
    st.markdown("## 📄 IEEE Report AI Agent (DeepSeek / Ollama)")

    # Ensure active project exists (set by app.py Dashboard)
    active_project = st.session_state.get("active_project")
    if not active_project:
        st.error("No active project found in session. Open this page from the main Dashboard.")
        return

    pid = safe_get(active_project, "project_id", "")
    project_title = safe_get(active_project, "project_name", "Untitled Project")
    domain = safe_get(active_project, "domain", "")

    st.info(f"**Project:** {project_title}  •  **Domain:** {domain}")

    # Init session state
    if "ieee_text" not in st.session_state:
        st.session_state.ieee_text = ""
    if "polished_text" not in st.session_state:
        st.session_state.polished_text = ""
    if "references_text" not in st.session_state:
        st.session_state.references_text = ""
    if "plag_score" not in st.session_state:
        st.session_state.plag_score = None
    if "eval_text" not in st.session_state:
        st.session_state.eval_text = ""

    # ---- Auto fetch from Mongo ----
    lit_text = fetch_project_literature(pid)
    obj_text = fetch_project_objectives(pid)
    impl_text = fetch_project_implementation(pid)

    lit_count = len(lit_text.splitlines()) if lit_text else 0
    obj_count = len(obj_text.splitlines()) if obj_text else 0
    task_count = len(impl_text.splitlines()) if impl_text else 0

    c1, c2, c3 = st.columns(3)
    c1.metric("Literature Lines", lit_count)
    c2.metric("Objectives", obj_count)
    c3.metric("Task/Impl Lines", task_count)

    last_report = get_last_report(pid)
    if last_report:
        created_at = last_report.get("created_at")
        auto_saved = last_report.get("auto_saved", False)
        st.caption(
            f"Last saved report version: {created_at} "
            f"({'auto' if auto_saved else 'manual'})"
        )

    with st.expander("📂 View Loaded Context", expanded=False):
        st.markdown("#### 🎯 Objectives (from DB)")
        st.write(obj_text or "_No objectives stored for this project._")

        st.markdown("#### 📘 Literature Review (from DB)")
        st.write(lit_text or "_No literature review stored for this project._")

        st.markdown("#### 🧩 Implementation / Tasks (from DB)")
        st.write(impl_text or "_No tasks / implementation stored for this project._")

    st.markdown("---")

    # ---- Generation options ----
    col_gen1, col_gen2 = st.columns(2)
    multi_chapter = col_gen1.checkbox(
        "Multi-chapter generation (slower, more detailed)", value=False
    )
    auto_save_toggle = col_gen2.checkbox(
        "Auto-save report to Mongo on generation", value=True
    )

    # ---- Generate IEEE Report ----
    if st.button("📝 Generate IEEE Report (AI Agent)"):
        if not project_title or not domain:
            st.error("Project title and domain must be set in the active project.")
        else:
            with st.spinner(
                "Generating IEEE report via DeepSeek (this may take a while)..."
            ):
                if multi_chapter:
                    text = generate_ieee_multi(
                        project_title, domain, obj_text, lit_text, impl_text
                    )
                else:
                    text = generate_ieee_single(
                        project_title, domain, obj_text, lit_text, impl_text
                    )

            if not text.strip():
                st.error(
                    "LLM returned empty text. Check Ollama server and model name "
                    f"('{OLLAMA_MODEL}')."
                )
            else:
                st.session_state.ieee_text = text
                st.session_state.polished_text = ""
                st.session_state.references_text = ""
                st.session_state.plag_score = None
                st.session_state.eval_text = ""
                if auto_save_toggle:
                    save_report_version(pid, project_title, domain, text, auto=True)
                st.success("IEEE report generated successfully!")

    st.markdown("### ✍️ Report Editor")

    base_text = st.session_state.ieee_text
    edited = st.text_area(
        "Report Text (you can manually edit here)",
        value=base_text,
        height=420,
    )
    st.session_state.ieee_text = edited

    # Manual Save Version
    if st.button("💾 Save Current Version to MongoDB"):
        if not edited.strip():
            st.warning("Nothing to save. Generate or paste a report first.")
        else:
            save_report_version(pid, project_title, domain, edited, auto=False)
            st.success("Report version saved to MongoDB (project_reports).")

    st.markdown("---")

    colA, colB, colC, colD = st.columns(4)

    # ---- Grammar polishing ----
    if colA.button("✨ Polish Grammar"):
        if not st.session_state.ieee_text.strip():
            st.warning("Nothing to polish. Generate or paste a report first.")
        else:
            with st.spinner("Polishing grammar and style via DeepSeek..."):
                polished = polish_text(st.session_state.ieee_text)
            if polished.strip():
                st.session_state.polished_text = polished
                st.success("Polished version generated below.")
            else:
                st.error("LLM failed to return polished text.")

    # ---- References ----
    if colB.button("📚 Generate References Only"):
        if not st.session_state.ieee_text.strip():
            st.warning("Provide report text first.")
        else:
            with st.spinner("Generating reference list from report..."):
                refs = generate_references(st.session_state.ieee_text, project_title, domain)
            if refs.strip():
                st.session_state.references_text = refs
                st.success("References generated below.")
            else:
                st.error("LLM failed to generate references.")

    # ---- Plagiarism (heuristic) ----
    if colC.button("🔍 Check Plagiarism (Heuristic)"):
        if not st.session_state.ieee_text.strip():
            st.warning("Provide report text first.")
        else:
            score = heuristic_plagiarism_score(st.session_state.ieee_text)
            st.session_state.plag_score = score
            st.info("Heuristic plagiarism risk score calculated (see below).")

    # ---- Clear ----
    if colD.button("🧹 Clear Session"):
        st.session_state.ieee_text = ""
        st.session_state.polished_text = ""
        st.session_state.references_text = ""
        st.session_state.plag_score = None
        st.session_state.eval_text = ""
        st.success("Cleared session values (report still in Mongo if saved).")

    # ---- Results: Polished + References + Plagiarism ----
    if st.session_state.polished_text:
        st.markdown("### ✨ Polished Version")
        st.text_area(
            "Polished Text (you can copy from here or replace main text manually)",
            value=str(st.session_state.polished_text),
            height=300,
        )

    if st.session_state.references_text:
        st.markdown("### 📚 Generated References")
        st.text_area(
            "References (paste into References section if needed)",
            value=str(st.session_state.references_text),
            height=220,
        )

    if st.session_state.plag_score is not None:
        st.markdown("### 🔍 Heuristic Plagiarism Score (Internal Only)")
        score = st.session_state.plag_score
        st.metric("Repetition Risk (0–100)", f"{score:.2f}")
        st.caption(
            "This is NOT a real plagiarism detector. "
            "It only checks internal repetition/self-similarity. "
            "It does NOT compare with internet / external sources."
        )
        st.write(explain_plagiarism_score(score))

    st.markdown("---")

    # ---- Evaluation & Suggestions ----
    st.markdown("### 🧠 AI Evaluation & Improvement Suggestions")

    if st.button("📊 Evaluate Report (AI Supervisor)"):
        if not st.session_state.ieee_text.strip():
            st.warning("Provide report text first.")
        else:
            with st.spinner("Evaluating IEEE structure, depth, and quality..."):
                eval_text = evaluate_report(st.session_state.ieee_text, project_title, domain)
            st.session_state.eval_text = eval_text or "Evaluation failed or empty."

    if st.session_state.eval_text:
        st.text_area(
            "Evaluation Output",
            value=str(st.session_state.eval_text),
            height=300,
        )

    st.markdown("---")

    # ---- Export Buttons ----
    st.markdown("### 📥 Export")

    final_text = st.session_state.ieee_text or ""
    if not final_text.strip():
        st.info("Generate or paste report text above to enable exports.")
        return

    # Text & Markdown download
    txt_bytes = final_text.encode("utf-8")
    base_name = project_title.replace(" ", "_") if project_title else "ieee_report"

    st.download_button(
        "⬇️ Download as .txt",
        data=txt_bytes,
        file_name=f"{base_name}.txt",
        mime="text/plain",
    )

    st.download_button(
        "⬇️ Download as .md",
        data=txt_bytes,
        file_name=f"{base_name}.md",
        mime="text/markdown",
    )

    # DOCX
    if Document is None:
        st.warning("python-docx not installed. DOCX export disabled.")
    else:
        docx_buf = export_docx(final_text, project_title)
        if docx_buf:
            st.download_button(
                "⬇️ Download as .docx",
                data=docx_buf,
                file_name=f"{base_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    # PDF
    if SimpleDocTemplate is None:
        st.warning("reportlab not installed. PDF export disabled.")
    else:
        pdf_buf = export_pdf(final_text, project_title)
        if pdf_buf:
            st.download_button(
                "⬇️ Download as .pdf",
                data=pdf_buf,
                file_name=f"{base_name}.pdf",
                mime="application/pdf",
            )


# When used as a Streamlit page inside TeamCollab, simply call:
ieee_report_agent_page()
